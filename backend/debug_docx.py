"""Debug script to analyze the XML structure of a DOCX file."""
import docx
from lxml import etree
import sys
import zipfile
import os

def debug_docx(path):
    print(f"=== Analyzing: {path} ===\n")
    
    # 1. Check what parts exist in the ZIP
    with zipfile.ZipFile(path, 'r') as z:
        names = z.namelist()
        print("=== ZIP contents (key files) ===")
        for n in names:
            if 'comment' in n.lower() or 'highlight' in n.lower() or '.xml' in n:
                print(f"  {n}")
        
        # Check if comments.xml exists
        if 'word/comments.xml' in names:
            print("\n=== word/comments.xml EXISTS ===")
            content = z.read('word/comments.xml').decode('utf-8')
            print(content[:2000])
        else:
            print("\n=== No word/comments.xml ===")
    
    # 2. Open with python-docx
    doc = docx.Document(path)
    
    # 3. Check for highlights
    highlights = doc.element.xpath('.//w:highlight')
    print(f"\n=== Highlights found: {len(highlights)} ===")
    for hl in highlights[:5]:
        print(f"  val={hl.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')}")
        # Show parent text
        run = hl.getparent().getparent()
        texts = run.xpath('.//w:t', namespaces=doc.element.nsmap)
        text = "".join(t.text for t in texts if t.text)
        print(f"  text: {text[:80]}")
    
    # 4. Check for comment references and ranges
    comment_refs = doc.element.xpath('.//w:commentReference')
    print(f"\n=== Comment References: {len(comment_refs)} ===")
    
    comment_range_starts = doc.element.xpath('.//w:commentRangeStart')
    print(f"=== Comment Range Starts: {len(comment_range_starts)} ===")
    
    comment_range_ends = doc.element.xpath('.//w:commentRangeEnd')
    print(f"=== Comment Range Ends: {len(comment_range_ends)} ===")
    
    # 5. Check for shading (background color)
    shd_elements = doc.element.xpath('.//w:rPr/w:shd')
    print(f"\n=== Run shading (w:rPr/w:shd): {len(shd_elements)} ===")
    
    pshd_elements = doc.element.xpath('.//w:pPr/w:shd')
    print(f"=== Paragraph shading (w:pPr/w:shd): {len(pshd_elements)} ===")
    
    # 6. Check for tables
    tables = doc.tables
    print(f"\n=== Tables found: {len(tables)} ===")
    for i, table in enumerate(tables):
        text = " ".join(cell.text for row in table.rows for cell in row.cells)
        print(f"  Table {i}: {text[:120]}...")
    
    # 7. Check for 'item' and 'score' in paragraphs
    print(f"\n=== Paragraphs with 'item' + 'score' or 'difficulty' ===")
    for p in doc.paragraphs:
        t = p.text.lower()
        if ('item' in t and 'score' in t) or 'difficulty' in t or 'proportion' in t or 'highest group' in t:
            print(f"  {p.text[:120]}")

if __name__ == "__main__":
    path = sys.argv[1] if len(sys.argv) > 1 else "/app/test.docx"
    debug_docx(path)
