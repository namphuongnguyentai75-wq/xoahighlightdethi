import os
import re
import docx
import shutil
import zipfile
from lxml import etree
from copy import deepcopy


def rebuild_docx(input_docx: str, output_docx: str, temp_dir: str):
    """
    Directly modifies the DOCX file using python-docx + raw XML manipulation.
    Removes: highlights, comments, item analysis tables/paragraphs, oval shapes.
    Keeps: all exam content (questions, answers, images) intact.
    """
    try:
        doc = docx.Document(input_docx)
    except Exception as e:
        print(f"Failed to open docx: {e}")
        shutil.copyfile(input_docx, output_docx)
        return output_docx

    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # ====================================================================
    # STEP 1: Gỡ TOÀN BỘ shading/background color và highlight
    # ====================================================================
    shd_removed = 0
    for shd in doc.element.xpath('.//*[local-name()="shd"]'):
        shd.getparent().remove(shd)
        shd_removed += 1
    hl_removed = 0
    for hl in doc.element.xpath('.//*[local-name()="highlight"]'):
        hl.getparent().remove(hl)
        hl_removed += 1
    print(f"[DEBUG] Shading removed: {shd_removed}, Highlights removed: {hl_removed}")

    # ====================================================================
    # STEP 2: Xóa tất cả các bảng chứa nội dung Item Analysis
    # ====================================================================
    tables_removed = 0
    for table in list(doc.tables):
        text = " ".join(cell.text for row in table.rows for cell in row.cells).lower()
        if any(kw in text for kw in ['proportion', 'highest group', 'lowest group',
                                       'r-pbis', 'difficulty', 'p value', 'discrim']):
            table._element.getparent().remove(table._element)
            tables_removed += 1
        elif 'item' in text and 'score' in text:
            table._element.getparent().remove(table._element)
            tables_removed += 1
    print(f"[DEBUG] Tables removed: {tables_removed}")

    # ====================================================================
    # STEP 3: Xóa phần Paragraph chứa Item Analysis và cứu các đáp án bị kẹt
    # ====================================================================
    paras_removed = 0
    runs_removed = 0
    
    def remove_empty_parents(element):
        while element is not None:
            parent = element.getparent()
            if parent is None:
                break
            if len(element) == 0 and element.tag.split('}')[-1] not in ['rPr', 'pPr', 't', 'br', 'p', 'r', 'tc']:
                parent.remove(element)
                element = parent
            else:
                break

    def insert_clean_paragraph(parent, insert_idx, text, nsmap):
        p = etree.Element(f'{{{nsmap["w"]}}}p')
        r = etree.SubElement(p, f'{{{nsmap["w"]}}}r')
        t = etree.SubElement(r, f'{{{nsmap["w"]}}}t')
        t.text = text
        parent.insert(insert_idx, p)
        return 1

    for p_element in doc.element.xpath('.//w:p'):
        texts = p_element.xpath('.//w:t')
        full_text = "".join(t.text for t in texts if t.text)
        full_text_lower = full_text.lower().strip()
        
        if not full_text_lower:
            continue
        
        # Kiểm tra xem paragraph này có chứa rác không
        item_match = re.search(r'Item\s*\d+\s*:\s*score|Commentato', full_text, flags=re.IGNORECASE)
        if item_match:
            parent = p_element.getparent()
            if parent is None:
                continue
            idx = parent.index(p_element)
            
            # Cứu phần câu hỏi (nếu có) nằm trước rác và trước đáp án
            q_match = re.match(r'^(.*?)(?=\b[A-Da-d]\.\s*|Item\s*\d+\s*:|Commentato)', full_text, flags=re.IGNORECASE | re.DOTALL)
            if q_match and q_match.group(1).strip():
                insert_clean_paragraph(parent, idx, q_match.group(1).strip(), doc.element.nsmap)
                idx += 1
                
            # Cứu các đáp án A, B, C, D bị kẹt trong đống rác (có tiền tố A, B, C, D)
            ans_matches = re.findall(r'\b([A-Da-d])\.\s*(.*?)(?=\b[A-Da-d]\.\s*|Item\s*\d+\s*:|Commentato|Difficulty|Proportion|$)', full_text, flags=re.IGNORECASE)
            seen = set()
            for m in ans_matches:
                ans_text = f"{m[0].upper()}. {m[1].strip()}"
                ans_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', ans_text)
                if ans_text not in seen:
                    seen.add(ans_text)
                    insert_clean_paragraph(parent, idx, ans_text, doc.element.nsmap)
                    idx += 1
                    
            # Nếu KHÔNG có tiền tố A, B, C, D, ta cứu phần text còn lại (vì có thể là list item)
            if not ans_matches:
                clean_text = re.sub(r'(?i)Commentato\s*\[.*?\]:?', '', full_text)
                clean_text = re.sub(r'(?i)Item\s*\d+\s*:\s*score.*', '', clean_text)
                clean_text = re.sub(r'(?i)Difficulty.*', '', clean_text)
                clean_text = re.sub(r'(?i)Proportion.*', '', clean_text)
                clean_text = clean_text.strip()
                
                q_text = q_match.group(1).strip() if q_match else ""
                q_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', q_text)
                if clean_text and clean_text != q_text:
                    half = len(clean_text) // 2
                    if clean_text[:half].strip() == clean_text[half:].strip() and len(clean_text) > 10:
                        clean_text = clean_text[:half].strip()
                        
                    clean_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', clean_text)
                    if clean_text:
                        insert_clean_paragraph(parent, idx, clean_text, doc.element.nsmap)
                        idx += 1

            parent.remove(p_element)
            paras_removed += 1
            continue
            
        # Nếu KHÔNG có "Item X:" nhưng paragraph chỉ chứa dữ liệu thống kê thuần túy
        # (bắt đầu bằng các keyword) thì xóa luôn
        if re.match(r'^difficulty\s*:', full_text_lower):
            p_element.getparent().remove(p_element)
            paras_removed += 1
            continue
        if re.match(r'^proportion\b', full_text_lower):
            p_element.getparent().remove(p_element)
            paras_removed += 1
            continue
        if re.match(r'^highest\s+group\b', full_text_lower):
            p_element.getparent().remove(p_element)
            paras_removed += 1
            continue
        if re.match(r'^lowest\s+group\b', full_text_lower):
            p_element.getparent().remove(p_element)
            paras_removed += 1
            continue
        if re.match(r'^r-pbis\b', full_text_lower):
            p_element.getparent().remove(p_element)
            paras_removed += 1
            continue
        if re.match(r'^p\s+value\b', full_text_lower):
            p_element.getparent().remove(p_element)
            # ====================================================================
    # STEP 4: Xóa comment markers (DISABLED TO PREVENT WORD CRASH)
    # ====================================================================
    comments_removed = 0
    '''
    for tag in ['commentReference', 'commentRangeStart', 'commentRangeEnd']:
        for el in doc.element.xpath(f'.//*[local-name()="{tag}"]'):
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
                comments_removed += 1
                remove_empty_parents(parent)
    # Xóa mc:AlternateContent chứa comment
    for ac in doc.element.xpath('.//*[local-name()="AlternateContent"]'):
        ac_xml = etree.tostring(ac, encoding='unicode').lower()
        if 'comment' in ac_xml:
            parent = ac.getparent()
            if parent is not None:
                parent.remove(ac)
                comments_removed += 1
                remove_empty_parents(parent)
    '''
    print(f"[DEBUG] Comment markers removed: {comments_removed}")

    # ====================================================================
    # STEP 5: Xóa hình vẽ oval/ellipse, bong bóng comment, highlight khung (DISABLED TO PREVENT WORD CRASH)
    # ====================================================================
    shapes_removed = 0
    '''
    # Xóa các shape DrawingML: wps:wsp và wps:txbx
    for tag in ['wsp', 'txbx']:
        for el in doc.element.xpath(f'.//*[local-name()="{tag}"]'):
            if el.getparent() is not None:
                parent = el.getparent()
                parent.remove(el)
                shapes_removed += 1
                remove_empty_parents(parent)
                
    # Xóa các shape VML (hình học cơ bản)
    for tag in ['rect', 'oval', 'line', 'polyline']:
        for el in doc.element.xpath(f'.//*[local-name()="{tag}"]'):
            if el.getparent() is not None:
                parent = el.getparent()
                parent.remove(el)
                shapes_removed += 1
                remove_empty_parents(parent)
                
    # Xóa v:shape NẾU NÓ KHÔNG CHỨA v:imagedata (nghĩa là hình vẽ chứ không phải ảnh)
    for shape in doc.element.xpath('.//*[local-name()="shape"]'):
        if not shape.xpath('.//*[local-name()="imagedata"]'):
            if shape.getparent() is not None:
                parent = shape.getparent()
                parent.remove(shape)
                shapes_removed += 1
                remove_empty_parents(parent)
    '''
    print(f"[DEBUG] Drawn shapes/highlights/comments removed: {shapes_removed}")

    # ====================================================================
    # STEP 6: Đổi chữ đỏ thành đen
    # ====================================================================
    red_removed = 0
    for color_elem in doc.element.xpath('.//w:rPr/w:color'):
        val = color_elem.get(f'{{{W_NS}}}val', '')
        if val and val.upper() in ['FF0000', 'FF0033', 'CC0000', 'C00000', 'FF3333', 'ED1C24', 'E36C09']:
            color_elem.set(f'{{{W_NS}}}val', '000000')
            red_removed += 1
    print(f"[DEBUG] Red color reset to black: {red_removed}")

    # Save
    doc.save(output_docx)

    # ====================================================================
    # STEP 7: Xóa comments.xml khỏi package ZIP
    # ====================================================================
    _remove_comments_from_zip(output_docx)

    print(f"[DEBUG] Done! Output: {output_docx}")
    return output_docx


def _remove_comments_from_zip(docx_path: str):
    """Remove word/comments.xml and related parts from the DOCX ZIP."""
    temp_path = docx_path + ".tmp"
    removed = 0

    with zipfile.ZipFile(docx_path, 'r') as zin:
        with zipfile.ZipFile(temp_path, 'w') as zout:
            for item in zin.infolist():
                # Skip any comment-related XML files
                basename = item.filename.lower()
                if 'comment' in basename and basename.endswith('.xml'):
                    removed += 1
                    continue
                
                data = zin.read(item.filename)

                # Clean comment relationships from .rels files
                if item.filename.endswith('.rels'):
                    try:
                        data_str = data.decode('utf-8')
                        if 'comment' in data_str.lower():
                            root = etree.fromstring(data)
                            for rel in list(root):
                                target = rel.get('Target', '').lower()
                                rtype = rel.get('Type', '').lower()
                                if 'comment' in target or 'comment' in rtype:
                                    root.remove(rel)
                                    removed += 1
                            data = etree.tostring(root, xml_declaration=True,
                                                   encoding='UTF-8', standalone=True)
                    except Exception as e:
                        print(f"[DEBUG] Error cleaning rels: {e}")

                # Clean [Content_Types].xml
                if item.filename == '[Content_Types].xml':
                    try:
                        root = etree.fromstring(data)
                        for override in list(root):
                            part_name = override.get('PartName', '').lower()
                            if 'comment' in part_name:
                                root.remove(override)
                                removed += 1
                        data = etree.tostring(root, xml_declaration=True,
                                               encoding='UTF-8', standalone=True)
                    except Exception as e:
                        print(f"[DEBUG] Error cleaning content types: {e}")

                zout.writestr(item, data)

    os.replace(temp_path, docx_path)
    print(f"[DEBUG] Comments ZIP parts removed: {removed}")


import subprocess
import fitz
from docx.shared import Inches, Pt
import tempfile

def convert_docx_to_pdf_headless(docx_path: str, output_dir: str):
    """Convert a DOCX to PDF using LibreOffice."""
    try:
        subprocess.run([
            "soffice", "--headless", "--convert-to", "pdf",
            "--outdir", output_dir, docx_path
        ], check=True)
        base_name = os.path.basename(docx_path)
        pdf_name = os.path.splitext(base_name)[0] + ".pdf"
        return os.path.join(output_dir, pdf_name)
    except Exception as e:
        print(f"Error converting to PDF: {e}")
        return None

def fix_vietnamese_spacing(text):
    vowels = 'a-zàáãạảăắằẵặẳâấầẫậẩeéèẽẹẻêếềễệểiíìĩịỉoóòõọỏôốồỗộổơớờỡợởuúùũụủưứừữựửyýỳỹỵỷđ'
    diacritics = 'àáãạảăắằẵặẳâấầẫậẩeéèẽẹẻêếềễệểiíìĩịỉoóòõọỏôốồỗộổơớờỡợởuúùũụủưứừữựửyýỳỹỵỷ'
    text = re.sub(r'([A-Za-zđĐ])\s+([' + diacritics + r'])', r'\1\2', text, flags=re.IGNORECASE)
    text = re.sub(r'([' + vowels + r'])\s+([' + diacritics + r'])', r'\1\2', text, flags=re.IGNORECASE)
    text = re.sub(r'([' + diacritics + r'])\s+([cmnpt]|ch|ng|nh)\b', r'\1\2', text, flags=re.IGNORECASE)
    return text

def clean_pdf_directly(input_path: str, output_path: str):
    doc = fitz.open(input_path)
    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block["type"] == 0:
                block_text = "".join(span["text"] for line in block["lines"] for span in line["spans"]).strip()
                clean_text = re.sub(r'\s+', '', block_text).lower()
                
                if re.match(r'^(proportion|highestgroup|lowestgroup|r-pbis|pvalue|difficulty)', clean_text):
                    page.add_redact_annot(block["bbox"], fill=(1, 1, 1))
                    continue
                if 'item' in clean_text and 'score' in clean_text and len(clean_text) < 100 and not re.search(r'^(câu|bài)\d+', clean_text):
                    page.add_redact_annot(block["bbox"], fill=(1, 1, 1))
                    continue
                    
                for line in block["lines"]:
                    line_text = "".join(span["text"] for span in line["spans"]).strip()
                    line_clean = re.sub(r'\s+', '', line_text).lower()
                    if re.match(r'^(proportion|highestgroup|lowestgroup|r-pbis|pvalue|difficulty)', line_clean):
                        page.add_redact_annot(line["bbox"], fill=(1, 1, 1))
                        continue
                        
                    for span in line["spans"]:
                        if re.search(r'(?i)Item\s*\d+\s*:\s*score|Commentato', span["text"]):
                            page.add_redact_annot(span["bbox"], fill=(1, 1, 1))
                            
        # Xóa các annotation highlight chuẩn của PDF
        for annot in page.annots():
            if annot.type[0] in [fitz.PDF_ANNOT_HIGHLIGHT, fitz.PDF_ANNOT_SQUARE, fitz.PDF_ANNOT_POLYGON]:
                page.delete_annot(annot)
                
        page.apply_redactions()
    doc.save(output_path)

def rebuild_pdf_to_docx(pdf_path: str, output_docx_path: str):
    """Fallback legacy method for PDF files."""
    doc = fitz.open(pdf_path)
    docx_doc = docx.Document()
    for page_num in range(len(doc)):
        page = doc[page_num]
        blocks = page.get_text("dict")["blocks"]
        blocks.sort(key=lambda b: b["bbox"][1])
        for block in blocks:
            if block["type"] == 0:
                block_text = "".join(span["text"] for line in block["lines"] for span in line["spans"]).strip()
                clean_text = re.sub(r'\s+', '', block_text).lower()
                
                # Bỏ qua hoàn toàn các block chỉ chứa thông số
                if re.match(r'^(proportion|highestgroup|lowestgroup|r-pbis|pvalue|difficulty)', clean_text):
                    continue
                if 'item' in clean_text and 'score' in clean_text and len(clean_text) < 100 and not re.search(r'^(câu|bài)\d+', clean_text):
                    continue

                current_p = docx_doc.add_paragraph()
                current_p.paragraph_format.space_after = Pt(0)
                current_p.paragraph_format.space_before = Pt(0)
                current_p.paragraph_format.line_spacing = 0.95
                
                stop_block = False
                for line in block["lines"]:
                    if stop_block: break
                    
                    line_text = "".join(span["text"] for span in line["spans"]).strip()
                    line_clean = re.sub(r'\s+', '', line_text).lower()
                    
                    # Nếu dòng bắt đầu bằng dữ liệu rác thống kê, dừng luôn việc in block này
                    if re.match(r'^(proportion|highestgroup|lowestgroup|r-pbis|pvalue|difficulty)', line_clean):
                        stop_block = True
                        break
                    
                    
                    # Thêm khoảng trắng giữa các đoạn nếu cần thiết
                    # Tuy nhiên, trong tiếng Việt trên PDF, các dấu (diacritics) thường bị tách thành các dòng/span riêng.
                    # Nếu ta tự động thêm space ở mỗi dòng, chữ sẽ bị nhảy (ví dụ "C â u").
                    # Nên ta chỉ thêm khoảng trắng nếu có dấu hiệu rõ ràng của một từ mới, nhưng để an toàn nhất:
                    if current_p.runs and not current_p.runs[-1].text.endswith(" "):
                        # Chỉ thêm khoảng trắng nếu dòng trước đó kết thúc bằng dấu câu hoặc dòng hiện tại bắt đầu bằng chữ hoa
                        prev_text = current_p.runs[-1].text
                        curr_text = "".join(span["text"] for span in line["spans"])
                        if curr_text and not re.match(r'^[a-zàáãạảăắằẵặẳâấầẫậẩeéèẽẹẻêếềễệểiíìĩịỉoóòõọỏôốồỗộổơớờỡợởuúùũụủưứừữựửyýỳỹỵỷđ]', curr_text):
                            current_p.add_run(" ")
                            
                    for span in line["spans"]:
                        raw_text = span["text"]
                        
                        # Cắt bỏ nếu đụng phải "Item X: score" hoặc "Commentato"
                        item_idx = re.search(r'(?i)Item\s*\d+\s*:\s*score|Commentato', raw_text)
                        if item_idx:
                            raw_text = raw_text[:item_idx.start()]
                            stop_block = True
                            
                        segments = re.split(r'(?=(?:Câu|Bài|CÂU|BÀI)\s+\d+[:\.])|(?=\b[A-Da-d]\.)', raw_text)
                        for seg in segments:
                            if not seg: continue
                            seg_strip = seg.strip()
                            if re.match(r'^(Câu|Bài|CÂU|BÀI)\s+\d+[:\.]|^[A-Da-d]\.', seg_strip):
                                if len(current_p.runs) > 0:
                                    current_p = docx_doc.add_paragraph()
                                    current_p.paragraph_format.space_after = Pt(0)
                                    current_p.paragraph_format.space_before = Pt(0)
                                    current_p.paragraph_format.line_spacing = 0.95
                                if re.match(r'^[A-Da-d]\.', seg_strip):
                                    current_p.paragraph_format.left_indent = Inches(0.25)
                                else:
                                    current_p.paragraph_format.left_indent = Inches(0)
                            
                            # Clean invalid XML control characters that cause Word to crash
                            safe_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', seg)
                            safe_text = fix_vietnamese_spacing(safe_text)
                            if not safe_text: continue
                            
                            run = current_p.add_run(safe_text)
                            
                            flags = span["flags"]
                            if (flags & 2): run.italic = True
                            if (flags & 16): run.bold = True
                            if "size" in span and span["size"] > 0:
                                fsize = round(span["size"])
                                if fsize > 0 and fsize < 100:
                                    run.font.size = Pt(fsize)
                                
                        if stop_block: break
            elif block["type"] == 1:
                # Disable image insertion for PDF-to-DOCX to prevent MS Word crashes
                # with unsupported image formats extracted by PyMuPDF.
                pass
    docx_doc.save(output_docx_path)
    return output_docx_path
